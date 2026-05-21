from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(r"C:\Users\User\Downloads\Manuela_Apresentacao_Institucional.docx")
OUT = Path(r"C:\Users\User\Downloads\Manuela_Apresentacao_Institucional_CONFERIDO_CODIGO.docx")


REPLACEMENTS = {
    "Defensoria Pública do Estado do Pará": "Rede de Proteção de Horizonte/CE",
    "Defensoria Pública do Pará": "rede de proteção de Horizonte/CE, com referência à Defensoria Pública do Estado do Ceará",
    "Defensoria Pública De Horizonte/CE": "rede de proteção de Horizonte/CE",
    "Delegacia da Mulher": "Delegacia Metropolitana de Horizonte",
    "A Defensoria Pública do Estado do Pará é o principal braço jurídico de proteção das pessoas em vulnerabilidade social no estado.": (
        "A rede de proteção de Horizonte/CE, com apoio da Defensoria Pública do Estado do Ceará, é o recorte territorial usado no MVP para orientar mulheres em situação de vulnerabilidade social."
    ),
    "A Rede de Proteção de Horizonte/CE é o principal braço jurídico de proteção das pessoas em vulnerabilidade social no estado.": (
        "A rede de proteção de Horizonte/CE, com apoio da Defensoria Pública do Estado do Ceará, é o recorte territorial usado no MVP para orientar mulheres em situação de vulnerabilidade social."
    ),
    "Frontend hospedado no Netlify com CORS configurado para o domínio da API": (
        "Frontend estático servido pelo Flask neste MVP; em produção pode ser separado em Netlify com CORS configurado para o domínio da API"
    ),
    "UptimeRobot realiza ping no endpoint /health a cada 5 minutos, mantendo o backend ativo": (
        "Endpoint /health disponível para monitoramento; o código também possui keep-alive interno a cada 10 minutos quando RENDER_EXTERNAL_URL está configurada"
    ),
    "Render.yaml configura disco persistente para manter o ChromaDB e o banco SQLite entre deploys": (
        "render.yaml configura o serviço web no Render; disco persistente deve ser configurado na plataforma antes de produção para preservar ChromaDB e SQLite entre deploys"
    ),
    "53 testes automatizados + 10 cenários críticos em bateria viva": "56 testes automatizados + 10 cenários críticos em bateria viva",
    "53 testes — todos passaram após as correções aplicadas": "56 testes — todos passaram após as correções aplicadas",
    "53 testes unitários + 10 cenários críticos em bateria viva": "56 testes unitários + 10 cenários críticos em bateria viva",
    "Frontend || HTML/CSS/JavaScript (Netlify) || Interface discreta “Dicas de Casa”; troca automática para modo protetivo": (
        "Frontend || HTML/CSS/JavaScript servido pelo Flask no MVP (Netlify opcional em produção) || Interface discreta “Dicas de Casa”; troca automática para modo protetivo"
    ),
    "HTML/CSS/JavaScript (Netlify)": "HTML/CSS/JavaScript servido pelo Flask no MVP (Netlify opcional em produção)",
    "Deploy || Netlify + Render + UptimeRobot || Frontend estático; backend persistente; ping a cada 5 min contra sleep": (
        "Deploy || Render + endpoint /health; Netlify/UptimeRobot opcionais se configurados || Frontend servido pelo Flask no MVP; backend web; persistência exige disco configurado no Render"
    ),
    "Netlify + Render + UptimeRobot": "Render + endpoint /health; Netlify/UptimeRobot opcionais se configurados",
    "Frontend estático; backend persistente; ping a cada 5 min contra sleep": (
        "Frontend servido pelo Flask no MVP; backend web; persistência exige disco configurado no Render"
    ),
    "Banco Vetorial || ChromaDB || Armazena chunks do manual jurídico; recupera contexto por similaridade": (
        "Banco Vetorial || ChromaDB || Armazena chunks do manual jurídico quando indexado; recupera contexto por similaridade"
    ),
    "Classificação || TF-IDF + Random Forest (scikit-learn) || Detecta sinais de risco; ~85–86% de acurácia; memória mínima (≤1 MB)": (
        "Classificação || TF-IDF + Random Forest (scikit-learn) || Classifica tipo e gravidade; F1 ponderado CV: tipo 0,8348 e gravidade 0,8548; modelos somam cerca de 20,7 MB"
    ),
    "Detecta sinais de risco; ~85–86% de acurácia; memória mínima (≤1 MB)": (
        "Classifica tipo e gravidade; F1 ponderado CV: tipo 0,8348 e gravidade 0,8548; modelos somam cerca de 20,7 MB"
    ),
    "5. RAG com ChromaDB || A mensagem é convertida em vetor (Gemini) e recupera os chunks mais relevantes do manual jurídico": (
        "5. RAG com ChromaDB || Quando a coleção está indexada, a mensagem é convertida em vetor (Gemini) e recupera chunks relevantes do manual jurídico"
    ),
    "Acurácia do classificador ML || ~85–86% em dataset sintético de 2.000 exemplos": (
        "Métrica do classificador ML || F1 ponderado em validação cruzada: tipo 0,8348; gravidade 0,8548"
    ),
    "Acurácia do classificador ML": "Métrica do classificador ML",
    "~85–86% em dataset sintético de 2.000 exemplos": "F1 ponderado em validação cruzada: tipo 0,8348; gravidade 0,8548",
    "Vocabulário treinado || Variações regionais brasileiras incluídas no dataset sintético": (
        "Dataset de treinamento || 1.880 exemplos sintéticos válidos, com 6 tipos de violência e 3 níveis de gravidade"
    ),
    "Delegacia da Mulher || Pedido de boletim de ocorrência || Registro de denúncia e inquérito": (
        "Delegacia Metropolitana de Horizonte || Risco, proteção física ou orientação policial local || Atendimento policial presencial conforme rede local mapeada no código"
    ),
    "Pedido de boletim de ocorrência": "Risco, proteção física ou orientação policial local",
    "Registro de denúncia e inquérito": "Atendimento policial presencial conforme rede local mapeada no código",
    "Medida Protetiva || Fluxo de denúncia, após BO || Proteção jurídica imediata — Lei Maria da Penha": (
        "Medida Protetiva || Pedido de denúncia ou orientação jurídica || Formulário oficial da Polícia Civil encaminhado ao Judiciário"
    ),
    "Casa da Mulher || Pedido de abrigo, falta de lugar seguro || Acolhimento físico, apoio psicossocial": (
        "Casa da Mulher Horizontina || Pedido de abrigo, falta de lugar seguro ou necessidade de acolhimento || Acolhimento e orientação; disponibilidade de abrigo deve ser confirmada pela rede humana"
    ),
    "Instituição || Defensoria Pública do Estado do Pará": (
        "Instituição || Rede de Proteção de Horizonte/CE / Defensoria Pública do Estado do Ceará"
    ),
    "Rede de Proteção de Horizonte/CE": "Rede de Proteção de Horizonte/CE / Defensoria Pública do Estado do Ceará",
    "Deploy || Netlify (frontend) + Render (backend)": (
        "Deploy || Render/Flask no MVP; Netlify opcional para frontend separado em produção"
    ),
    "Netlify (frontend) + Render (backend)": "Render/Flask no MVP; Netlify opcional para frontend separado em produção",
    "Classificação || ~85–86% acurácia (TF-IDF + Random Forest, dataset 2.000 exemplos)": (
        "Classificação || F1 ponderado CV: tipo 0,8348; gravidade 0,8548 (TF-IDF + Random Forest, 1.880 exemplos válidos)"
    ),
    "~85–86% acurácia (TF-IDF + Random Forest, dataset 2.000 exemplos)": (
        "F1 ponderado CV: tipo 0,8348; gravidade 0,8548 (TF-IDF + Random Forest, 1.880 exemplos válidos)"
    ),
}


def replace_in_paragraph(paragraph, replacements):
    text = paragraph.text
    new = text
    for old, replacement in replacements.items():
        new = new.replace(old, replacement)
    if new == text:
        return
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new
    else:
        paragraph.add_run(new)


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    doc = Document(SRC)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, REPLACEMENTS)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, REPLACEMENTS)

    # Safety metadata: avoid personal author metadata in the deliverable copy.
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = "Conferido contra o código local em 20/05/2026."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
