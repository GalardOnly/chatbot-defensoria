from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


SRC = Path(r"C:\Users\User\Downloads\Manuela_Apresentacao_Institucional.docx")
OUT = Path(r"C:\Users\User\Downloads\Manuela_Apresentacao_Institucional_ENTREGA_DEFENSORIA_PREFEITURA.docx")


REPLACEMENTS = {
    "Defensoria Pública do Estado do Pará": "Rede de Proteção de Horizonte/CE / Defensoria Pública do Estado do Ceará",
    "Defensoria Pública do Pará": "rede de proteção de Horizonte/CE, com referência à Defensoria Pública do Estado do Ceará",
    "Defensoria Pública De Horizonte/CE": "rede de proteção de Horizonte/CE",
    "Delegacia da Mulher": "Delegacia Metropolitana de Horizonte",
    "53 testes automatizados + 10 cenários críticos em bateria viva": "56 testes automatizados + 10 cenários críticos em bateria viva",
    "53 testes — todos passaram após as correções aplicadas": "56 testes — todos passaram após as correções aplicadas",
    "53 testes unitários + 10 cenários críticos em bateria viva": "56 testes unitários + 10 cenários críticos em bateria viva",
    "Netlify + Render + UptimeRobot": "Render + endpoint /health; Netlify/UptimeRobot opcionais se configurados",
    "Frontend estático; backend persistente; ping a cada 5 min contra sleep": "Frontend servido pelo Flask no MVP; backend web; persistência exige disco configurado no Render",
    "HTML/CSS/JavaScript (Netlify)": "HTML/CSS/JavaScript servido pelo Flask no MVP (Netlify opcional em produção)",
    "Detecta sinais de risco; ~85–86% de acurácia; memória mínima (≤1 MB)": "Classifica tipo e gravidade; F1 ponderado CV: tipo 0,8348 e gravidade 0,8548; modelos somam cerca de 20,7 MB",
    "Acurácia do classificador ML": "Métrica do classificador ML",
    "~85–86% em dataset sintético de 2.000 exemplos": "F1 ponderado em validação cruzada: tipo 0,8348; gravidade 0,8548",
    "Variações regionais brasileiras incluídas no dataset sintético": "1.880 exemplos sintéticos válidos, com 6 tipos de violência e 3 níveis de gravidade",
    "Pedido de boletim de ocorrência": "Risco, proteção física ou orientação policial local",
    "Registro de denúncia e inquérito": "Atendimento policial presencial conforme rede local mapeada no código",
    "Netlify (frontend) + Render (backend)": "Render/Flask no MVP; Netlify opcional para frontend separado em produção",
    "~85–86% acurácia (TF-IDF + Random Forest, dataset 2.000 exemplos)": "F1 ponderado CV: tipo 0,8348; gravidade 0,8548 (TF-IDF + Random Forest, 1.880 exemplos válidos)",
    "Render.yaml configura disco persistente para manter o ChromaDB e o banco SQLite entre deploys": "render.yaml configura o serviço web no Render; disco persistente deve ser configurado na plataforma antes de produção para preservar ChromaDB e SQLite entre deploys",
    "UptimeRobot realiza ping no endpoint /health a cada 5 minutos, mantendo o backend ativo": "Endpoint /health disponível para monitoramento; o código também possui keep-alive interno a cada 10 minutos quando RENDER_EXTERNAL_URL está configurada",
    "Frontend hospedado no Netlify com CORS configurado para o domínio da API": "Frontend estático servido pelo Flask neste MVP; em produção pode ser separado em Netlify com CORS configurado para o domínio da API",
}


SECTION_TEXT = [
    ("10. Apresentação à Defensoria ou Prefeitura", "Heading 1"),
    (
        "O objetivo recomendado para a primeira reunião não é pedir implantação imediata, mas solicitar avaliação institucional do MVP. A apresentação deve posicionar a Manuela como protótipo de apoio inicial, sujeito à validação jurídica, psicossocial, técnica e de governança de dados.",
        None,
    ),
    ("10.1 Objetivo da reunião", "Heading 2"),
    ("Validar se o problema atendido pelo MVP é relevante para a Defensoria, Prefeitura ou rede municipal de proteção.", "List Bullet"),
    ("Obter indicação de um ponto focal institucional para revisar linguagem, fluxos críticos, contatos oficiais e limites de atuação.", "List Bullet"),
    ("Definir se há interesse em um piloto controlado, sem uso público amplo e sem promessa de atendimento automatizado de emergência.", "List Bullet"),
    ("Levantar requisitos de LGPD, retenção de dados, acesso administrativo, auditoria, atendimento humano e encaminhamento de casos críticos.", "List Bullet"),
    ("10.2 Como conduzir a apresentação", "Heading 2"),
    ("Abrir com o aviso de limitação: apoio inicial, não substitui atendimento humano, jurídico, psicológico, policial ou médico.", "List Bullet"),
    ("Demonstrar somente mensagens sintéticas: saudação, 'nao posso falar', 'fui agredida', 'quero denunciar' e 'nao tenho para onde ir'.", "List Bullet"),
    ("Mostrar os controles de segurança: fachada visual discreta, botão Sair, botão Apagar, conversa anônima e logs sem conteúdo sensível.", "List Bullet"),
    ("Explicar que os contatos oficiais precisam ser validados pela instituição antes de qualquer uso real.", "List Bullet"),
    ("Encerrar pedindo avaliação formal, não autorização informal para produção.", "List Bullet"),
    ("10.3 Materiais sugeridos para levar", "Heading 2"),
    ("Documento institucional em PDF/DOCX, roteiro de demonstração, lista de testes críticos, README técnico e .env.example sem chaves reais.", "List Bullet"),
    ("Tela local do MVP já aberta, com servidor rodando, sem histórico real no banco e sem exibição de valores do arquivo .env.", "List Bullet"),
    ("Resumo de riscos residuais: necessidade de governança humana, política LGPD, revisão de conteúdo e infraestrutura de produção.", "List Bullet"),
    ("11. Próximos passos se a avaliação for positiva", "Heading 1"),
    (
        "Se a Defensoria, Prefeitura ou rede de proteção considerar o MVP promissor, a implementação deve avançar por fases, com pontos de decisão claros e supervisão humana. A recomendação é não publicar para uso aberto antes de um piloto controlado.",
        None,
    ),
    ("Fase 1 — Validação institucional (2 a 4 semanas)", "Heading 2"),
    ("Nomear responsáveis técnicos, jurídicos, psicossociais e de proteção de dados.", "List Bullet"),
    ("Revisar frases de risco, linguagem acolhedora, contatos oficiais, fluxos de denúncia, abrigo e medida protetiva.", "List Bullet"),
    ("Definir quais dados podem ser armazenados, por quanto tempo, quem acessa e como a usuária pode solicitar exclusão.", "List Bullet"),
    ("Fase 2 — Piloto controlado (4 a 8 semanas)", "Heading 2"),
    ("Implantar em ambiente restrito, com domínio controlado, HTTPS, CORS fechado e base de conhecimento revisada.", "List Bullet"),
    ("Usar apenas grupo pequeno de avaliadoras, servidoras ou equipe parceira; não divulgar ao público geral.", "List Bullet"),
    ("Registrar métricas sem conteúdo sensível: disponibilidade, falhas, tipo de encaminhamento e necessidade de intervenção humana.", "List Bullet"),
    ("Fase 3 — Ajustes e decisão de continuidade", "Heading 2"),
    ("Revisar respostas problemáticas encontradas no piloto e atualizar testes de regressão.", "List Bullet"),
    ("Confirmar se o fluxo reduz risco ou se cria dúvidas, falsa segurança ou carga indevida para a rede humana.", "List Bullet"),
    ("Decidir entre continuar, limitar escopo, transformar em ferramenta interna ou suspender a iniciativa.", "List Bullet"),
    ("Fase 4 — Implantação assistida", "Heading 2"),
    ("Configurar infraestrutura com backup seguro, rotação de chaves, monitoramento sem conteúdo privado e plano de resposta a incidentes.", "List Bullet"),
    ("Criar protocolo de atualização periódica dos contatos oficiais e responsável institucional por cada revisão.", "List Bullet"),
    ("Treinar equipe humana para interpretar encaminhamentos e manter canal de feedback contínuo.", "List Bullet"),
    ("Critérios mínimos para seguir adiante", "Heading 2"),
    ("A instituição aceita formalmente o escopo de apoio inicial e as limitações do chatbot.", "List Bullet"),
    ("Há responsável humano por revisão, incidentes, atualização de conteúdo e solicitações de exclusão.", "List Bullet"),
    ("Os testes críticos continuam passando após cada alteração.", "List Bullet"),
    ("Nenhum fluxo orienta confronto, fuga sem plano, exposição da usuária ou apagamento indevido de provas.", "List Bullet"),
]


def replace_in_paragraph(paragraph, replacements):
    text = paragraph.text
    new = text
    for old, replacement in replacements.items():
        new = new.replace(old, replacement)
    if new == text:
        return
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = new
    else:
        paragraph.add_run(new)


def set_inserted_style(paragraph, style_name):
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            if style_name == "List Bullet" and paragraph.runs:
                paragraph.runs[0].text = f"- {paragraph.runs[0].text}"
    if style_name == "Heading 1":
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(16)
            run.bold = True
    elif style_name == "Heading 2":
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(13)
            run.bold = True


def insert_sections_before_project_info(doc):
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("9. Informações do Projeto"):
            target = paragraph
            break
    if target is None:
        target = doc.paragraphs[-1]

    for text, style in SECTION_TEXT:
        p = target.insert_paragraph_before(text)
        set_inserted_style(p, style)


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    doc = Document(SRC)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, REPLACEMENTS)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, REPLACEMENTS)

    insert_sections_before_project_info(doc)

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = "Conferido contra o código local e ampliado com próximos passos institucionais."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
